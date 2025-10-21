"""
ELABORADOR DE PREGUNTAS AIKEN DE DOCX
"""

#!/usr/bin/env python
# coding: utf-8

# In[ ]:





# In[4]:


import os
from docx import Document
from openai import OpenAI
import time

def create_openai_client(api_key):
    return OpenAI(api_key=api_key)

def create_assistant(client, name, instructions, model="gpt-4o"):
    return client.beta.assistants.create(
        name=name,
        instructions=instructions,
        model=model,
        tools=[{"type": "file_search"}],
    )

def create_vector_store(client, name):
    return client.beta.vector_stores.create(name=name)

def extract_text_from_docx(file_paths):
    extracted_texts = []
    for path in file_paths:
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        extracted_texts.append(text)
    return extracted_texts

def save_texts_to_files(extracted_texts, output_dir, original_file_names):
    temp_txt_paths = []
    for i, text in enumerate(extracted_texts):
        # Use the original DOCX file name for the TXT file
        original_file_name = os.path.splitext(os.path.basename(original_file_names[i]))[0]
        temp_txt_path = os.path.join(output_dir, f"{original_file_name}.txt")
        with open(temp_txt_path, "w", encoding="utf-8") as temp_file:
            temp_file.write(text)
        temp_txt_paths.append(temp_txt_path)
    return temp_txt_paths

def upload_files_to_vector_store(client, vector_store_id, temp_txt_paths):
    file_streams = [open(path, "rb") for path in temp_txt_paths]
    file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
        vector_store_id=vector_store_id, files=file_streams
    )
    return file_batch

def create_message_file(client, file_path):
    return client.files.create(
        file=open(file_path, "rb"), purpose="assistants"
    )

def create_thread(client, assistant_id, message_file_id, content):
    thread = client.beta.threads.create(
        messages=[
            {
                "role": "user",
                "content": content,
                "attachments": [
                    {"file_id": message_file_id, "tools": [{"type": "file_search"}]}
                ],
            }
        ]
    )
    return thread

def run_thread(client, thread_id, assistant_id):
    # Intentamos obtener mensajes hasta un máximo de 5 veces con un retraso entre cada intento
    for attempt in range(5):
        run = client.beta.threads.runs.create_and_poll(
            thread_id=thread_id, assistant_id=assistant_id
        )
        messages = list(client.beta.threads.messages.list(thread_id=thread_id, run_id=run.id))

        if messages:
            return messages
        else:
            print(f"No messages returned from the thread, attempt {attempt + 1}/5. Retrying...")
            time.sleep(10)  # Espera 10 segundos antes de intentar nuevamente

    return None

def main():
    api_key = "YOUR_API_KEY_HERE"
    client = create_openai_client(api_key)

    assistant = create_assistant(
        client,
        name="Elaborador",
        instructions="""Eres un experto en elaboración de preguntas didácticas para evaluaciones. 
        No proporcionas las fuentes ni haces citaciones. 
        No das un texto de introducción, ni de despedida. 
        Únicamente respondes con las preguntas solicitadas en el formato AIKEN, sin más ni menos."""
    )

    vector_store = create_vector_store(client, name="Texto_adjunto")

    # Ruta del directorio donde están los archivos DOCX
    input_dir = r"C:\Users\HP\Desktop\CATO-CURSOS-2-2024\GER-TI CATO1-2024\Cursos\SEMANA 2\BUSINESS MODEL CANVAS\texto_videos"

    # Crear carpeta "preguntas" si no existe
    output_dir = os.path.join(input_dir, 'preguntas')
    os.makedirs(output_dir, exist_ok=True)

    # Lista de todos los archivos DOCX en el directorio
    file_paths = [os.path.join(input_dir, f) for f in os.listdir(input_dir) if f.endswith('.docx')]

    extracted_texts = extract_text_from_docx(file_paths)
    temp_txt_paths = save_texts_to_files(extracted_texts, output_dir, file_paths)

    file_batch = upload_files_to_vector_store(client, vector_store.id, temp_txt_paths)
    print("File batch status:", file_batch.status)
    print("File batch counts:", file_batch.file_counts)

    assistant = client.beta.assistants.update(
        assistant_id=assistant.id,
        tool_resources={"file_search": {"vector_store_ids": [vector_store.id]}},
    )

    for temp_txt_path in temp_txt_paths:
        message_file = create_message_file(client, temp_txt_path)

        content = """Elabora VEINTE preguntas en formato AIKEN basadas en el texto proporcionado. 
        Las preguntas deben centrarse en el tema principal del texto. 
        No debes proporcionar citaciones ni incluir introducción o despedida. 
        Solo responde con las preguntas en el formato AIKEN.

        Ejemplo de preguntas:

        Según el texto de título: Business Storytelling Masterclass with Matteo Cassese, ¿qué es esencial para captar la atención del público en storytelling?
        A) Utilizar terminología complicada
        B) Hablar en un tono monótono
        C) Empezar con una anécdota interesante
        D) Presentar gráficos complejos
        ANSWER: C
        """

        thread = create_thread(client, assistant.id, message_file.id, content)

        # Esperar a que el hilo se ejecute completamente
        messages = run_thread(client, thread.id, assistant.id)

        if messages:
            message_content = messages[0].content[0].text
            # Guardar las preguntas en un archivo TXT en la carpeta "preguntas"
            output_txt_path = os.path.join(output_dir, os.path.basename(temp_txt_path))
            with open(output_txt_path, "w", encoding="utf-8") as output_file:
                output_file.write(message_content.value)
            print(f"Preguntas guardadas en: {output_txt_path}")
        else:
            print("No messages returned from the thread after multiple attempts.")

        # Esperar un breve momento antes de pasar al siguiente archivo para asegurar que todo se haya completado
        time.sleep(2)  # Puedes ajustar el tiempo de espera si es necesario

if __name__ == "__main__":
    main()


# In[ ]:


import os
from docx import Document
from openai import OpenAI
import time

def create_openai_client(api_key):
    return OpenAI(api_key=api_key)

def create_assistant(client, name, instructions, model="gpt-4"):
    return client.beta.assistants.create(
        name=name,
        instructions=instructions,
        model=model,
        tools=[]
    )

def create_vector_store(client, name):
    return client.beta.vector_stores.create(name=name)

def extract_text_from_docx(file_paths):
    extracted_texts = []
    for path in file_paths:
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        extracted_texts.append(text)
    return extracted_texts

def save_texts_to_files(extracted_texts, output_dir, original_file_names):
    temp_txt_paths = []
    for i, text in enumerate(extracted_texts):
        original_file_name = os.path.splitext(os.path.basename(original_file_names[i]))[0]
        temp_txt_path = os.path.join(output_dir, f"{original_file_name}.txt")
        with open(temp_txt_path, "w", encoding="utf-8") as temp_file:
            temp_file.write(text)
        temp_txt_paths.append(temp_txt_path)
    return temp_txt_paths

def upload_files_to_vector_store(client, vector_store_id, temp_txt_paths):
    file_streams = [open(path, "rb") for path in temp_txt_paths]
    file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
        vector_store_id=vector_store_id, files=file_streams
    )
    return file_batch

def create_message_file(client, file_path):
    return client.files.create(
        file=open(file_path, "rb"), purpose="assistants"
    )

def create_thread(client, assistant_id, message_file_id, content):
    thread = client.beta.threads.create(
        messages=[
            {
                "role": "user",
                "content": content,
                "attachments": [
                    {"file_id": message_file_id}
                ],
            }
        ]
    )
    return thread

def run_thread(client, thread_id, assistant_id):
    for attempt in range(5):
        try:
            run = client.beta.threads.runs.create_and_poll(
                thread_id=thread_id, assistant_id=assistant_id
            )
            messages = list(client.beta.threads.messages.list(thread_id=thread_id, run_id=run.id))
            if messages:
                return messages
            else:
                print(f"No messages returned from the thread, attempt {attempt + 1}/5. Retrying...")
                time.sleep(10)
        except Exception as e:
            print(f"Error during thread run: {str(e)}. Attempt {attempt + 1}/5. Retrying...")
            time.sleep(10)

    return None

def main():
    api_key = "YOUR_API_KEY_HERE"
    client = create_openai_client(api_key)

    assistant = create_assistant(
        client,
        name="Elaborador",
        instructions="""Eres un experto en elaboración de preguntas didácticas para evaluaciones. 
        No proporcionas las fuentes ni haces citaciones. 
        No das un texto de introducción, ni de despedida. 
        Únicamente respondes con las preguntas solicitadas en el formato AIKEN, sin más ni menos."""
    )

    vector_store = create_vector_store(client, name="Texto_adjunto")

    input_dir = r"C:\Users\HP\Desktop\CATO-CURSOS-2-2024\GER-TI CATO1-2024\Cursos\SEMANA 2\BUSINESS MODEL CANVAS\texto_videos"
    output_dir = os.path.join(input_dir, 'preguntas')
    os.makedirs(output_dir, exist_ok=True)

    file_paths = [os.path.join(input_dir, f) for f in os.listdir(input_dir) if f.endswith('.docx')]

    extracted_texts = extract_text_from_docx(file_paths)
    temp_txt_paths = save_texts_to_files(extracted_texts, output_dir, file_paths)

    file_batch = upload_files_to_vector_store(client, vector_store.id, temp_txt_paths)
    print("File batch status:", file_batch.status)
    print("File batch counts:", file_batch.file_counts)

    assistant = client.beta.assistants.update(
        assistant_id=assistant.id,
        tool_resources={"file_search": {"vector_store_ids": [vector_store.id]}},
    )

    for temp_txt_path in temp_txt_paths:
        message_file = create_message_file(client, temp_txt_path)

        content = """Elabora VEINTE preguntas en formato AIKEN basadas en el texto proporcionado. 
        Las preguntas deben centrarse en el tema principal del texto. 
        No debes proporcionar citaciones ni incluir introducción o despedida. 
        Solo responde con las preguntas en el formato AIKEN."""

        thread = create_thread(client, assistant.id, message_file.id, content)

        messages = run_thread(client, thread.id, assistant.id)

        if messages:
            message_content = messages[0].content
            output_txt_path = os.path.join(output_dir, os.path.basename(temp_txt_path))
            with open(output_txt_path, "w", encoding="utf-8") as output_file:
                output_file.write(message_content)
            print(f"Preguntas guardadas en: {output_txt_path}")
        else:
            print("No messages returned from the thread after multiple attempts.")

        time.sleep(2)

if __name__ == "__main__":
    main()


# In[ ]:




